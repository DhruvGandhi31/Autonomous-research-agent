import io
from pathlib import Path

from loguru import logger


class DocumentExtractor:
    """Extracts plain text from PDF and DOCX files."""

    def extract(self, content: bytes, filename: str) -> dict:
        suffix = Path(filename).suffix.lower()
        try:
            if suffix == ".pdf":
                return self._extract_pdf(content)
            elif suffix in (".docx", ".doc"):
                return self._extract_docx(content)
            elif suffix == ".txt":
                text = content.decode("utf-8", errors="replace")
                return {"text": text, "pages": 1, "method": "plain_text"}
            else:
                return {"text": "", "pages": 0, "error": f"Unsupported format: {suffix}"}
        except Exception as e:
            logger.error(f"Document extraction failed for {filename}: {e}")
            return {"text": "", "pages": 0, "error": str(e)}

    def _extract_pdf(self, content: bytes) -> dict:
        import PyPDF2

        reader = PyPDF2.PdfReader(io.BytesIO(content))
        pages_text = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages_text.append(text)
        full_text = "\n\n".join(t for t in pages_text if t.strip())
        return {"text": full_text, "pages": len(reader.pages), "method": "pypdf2"}

    def _extract_docx(self, content: bytes) -> dict:
        from docx import Document

        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        full_text = "\n\n".join(paragraphs)
        return {"text": full_text, "pages": len(doc.paragraphs), "method": "python-docx"}


document_extractor = DocumentExtractor()
